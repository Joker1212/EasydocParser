#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#
import json
import os

from PIL import Image
from docx import Document
import re

from docx.oxml import CT_Tbl, CT_P, parse_xml
from docx.table import Table
from docx.text.paragraph import Paragraph

from io import BytesIO
from Levenshtein import ratio as levenshtein_ratio


from utils.file_util import get_output_path, get_output_dir
from utils.output_helper import LoaderTable, LoaderImage, LoaderOutPutEncoder, LoaderText
from utils.title_helper import is_possible_title



class DocxLoader:
    def __init__(self, **kwargs):
        super().__init__()
        self.filename = kwargs.get('filename')
        self.file = kwargs.get('file')
        self.debug = kwargs.get('debug', False)
        if not self.filename and not self.file:
            raise ValueError('filename or file must be provided')
        self.doc = Document(self.filename) if not self.file else Document(self.file)
        self.output_dir = kwargs.get('output_dir', get_output_dir(self.filename, output_dir=f'output'))
        self.toc_entries = []  # 用于存储目录项
        self.last_matched_index = -1  # 上次匹配成功的toc条目的索引
        self.page_number = 0  # 当前页数
        self.title_stack = []  # 存储标题堆栈
        self.level_stack = []  # 存储标题级别堆栈
        self.output = []  # 存储处理后的段落、表格、图片和标题信息
        self.min_img_height = 80
        self.min_img_width = 80
        self.image_num = 0
        self.table_num = 0
        self.page_number = 0
        self.title_stack, self.level_stack = [], []
        self.output = []
        # 用于避免图片重复以及表格中图片错误抽取
        self.img_ids = set()
        # 从toc中收集标题时，标题的默认最大长度，实际会根据toc的真实长度更新
        self.toc_max_length = 20
        self.has_title_xml = False



    def docx_title_find(self, p):
        clean_text = re.sub(r"\u3000", " ", p.text).strip()
        if p.style.name.startswith('Heading'):
            self.has_title_xml = True
            return int(p.style.name.split(' ')[-1]), clean_text, 'strong', True
        toc_header = self.get_toc_title(clean_text)
        if toc_header:
            self.has_title_xml = True
            return toc_header['level'], toc_header['header'], 'strong', True
        # # 有前置的toc标题，使用toc目录格式
        if self.has_title_xml:
            return 0, clean_text, None, False
        # # 参考langchian-chat-chat,既没有heading又没有toc目录，只能规则匹配
        if is_possible_title(clean_text):
            if len(self.level_stack) > 0:
                return self.level_stack[-1], clean_text, 'weak', True
            return 1, clean_text, 'weak', True
        return 0, clean_text, 'weak', False

    def clean_text(self, text):
        """清洗文本，去除特殊符号和前缀"""
        # 去除特殊符号
        # clean_text = re.sub(r'[^\w\s]', '', text)
        clean_text = re.sub(r'[\t\n].*$', '', text).strip()
        clean_text = re.sub(r"\u3000", " ", clean_text).strip()

        # 去除前缀
        cleaned_text = clean_text.lstrip("toc ")
        return cleaned_text.strip()

    def is_heading_style(self, style_name):
        """判断样式名称是否包含 "Heading" 标识"""
        return "Heading" in style_name

    def is_toc_entry(self, style_name):
        """判断样式名称是否包含 "toc" 标识"""
        return "toc" in style_name.lower()

    def get_toc_title(self, text, similarity_threshold=0.8):
        """判断文本是否与之前的toc文本语义基本一致"""
        # 清洗文本
        cleaned_text = self.clean_text(text)
        if len(cleaned_text) > self.toc_max_length:
            return None
        # 只与上次匹配成功的toc条目之后的所有条目进行相似度计算
        for toc_entry in self.toc_entries[self.last_matched_index + 1:]:
            similarity = levenshtein_ratio(cleaned_text, toc_entry['header'])
            if similarity >= similarity_threshold:
                self.last_matched_index += 1  # 更新上次匹配成功的索引
                return toc_entry
        return None

    def get_paragraph_info(self, text, page_number, title_stack, level_stack, filename) -> LoaderText:
        """获取段落信息"""
        return LoaderText(page_number, filename, title_stack, text, level_stack)

    def process_paragraph_with_image(self, paragraph, last_paragraph_text, output_list, page_number, title_stack,
                                     level_stack, filename):
        """处理包含图片的段落，返回图片信息列表"""
        """获取指定段落周围的文本信息，包括前向、后向和本段落的文本"""
        current_text = last_paragraph_text + paragraph.text.strip()
        # forward_text = self.get_forward_text(last_paragraph_text, output_list, title_stack)
        # surrounding_text = {'forward_text': forward_text, 'backward_text': None, 'current_text': current_text}
        images_info = []
        for run in paragraph.runs:
            if run._r.xpath('.//pic:pic'):  # 检查run中是否有图片
                for img in run._r.xpath('.//pic:pic'):
                    # 获取图片信息
                    image_info = self.get_image_info(img, page_number, title_stack, level_stack, filename, None,
                                                     f"isolated_{page_number}")
                    if image_info:
                        images_info.append(image_info)
        return (bool(current_text), images_info)  # 返回标志和图片信息列表

    def get_forward_text(self, last_paragraph_text, output_list, title_stack):
        forward_text = ''
        # if not len(forward_text):
        for i in range(len(output_list) - 1, -1, -1):
            if not output_list[i].get('text'):
                continue
            if all(item in title_stack for item in output_list[i]['title_stack']) and len(
                    output_list[i]['title_stack']) >= len(title_stack) - 1:
                forward_text = output_list[i]['text']
                break
        return forward_text

    def save_image(self, image_data, image_path):
        """保存图片到本地，并返回图片的URL"""

        # 生成图片文件名
        # image_filename = f"{filename_prefix}_{len(os.listdir(output_dir))}.png"
        # image_path = os.path.join(output_dir, image_filename)
        image_data.save(image_path)
        return image_path

    def process_table_with_context(self, table, last_paragraph_text, output_list, page_number, title_stack, level_stack,
                                   doc_filename) -> LoaderTable:
        """处理表格，并返回表格信息及前后文的文本信息"""
        forward_text = ''
        # forward_text = self.get_forward_text(last_paragraph_text, output_list, title_stack)
        surrounding_text = {'forward_text': forward_text, 'backward_text': None, 'current_text': last_paragraph_text}

        html_table = "<!DOCTYPE html>\n<html>\n<head>\n<meta charset='UTF-8'>\n</head>\n<body>\n<table>\n"
        headers = []
        rows = []
        pre_output = output_list[-1]

        # 遍历表格的每一行
        for row_index, row in enumerate(table.rows):
            # 处理表头
            if not headers:
                html_table += "  <thead>\n    <tr>\n"
                for cell_index, cell in enumerate(row.cells):
                    cell_text = self.get_cell_text(cell, doc_filename, level_stack, page_number, title_stack)
                    # 添加样式使表头居中对齐
                    html_table += f"      <th style='text-align: left;'>{cell_text.strip()}</th>\n"
                    headers.append(cell_text.strip())  # 添加表头到 headers 列表
                html_table += "    </tr>\n  </thead>\n  <tbody>\n"
                continue
            # 有重复的表格，直接返回 None
            if pre_output and isinstance(pre_output, LoaderTable) and all(
                    header in headers for header in pre_output.table_headers):
                return None
            # 处理数据行
            row_data = []
            html_table += f"    <tr data-row='{row_index}'>\n"
            for cell_index, cell in enumerate(row.cells):
                cell_text = self.get_cell_text(cell, doc_filename, level_stack, page_number, title_stack)
                # 添加样式使单元格内容居中对齐
                html_table += f"      <td style='text-align: left;' data-col='{cell_index}'>{cell_text}</td>\n"
            html_table += "    </tr>\n"
            rows.append(row_data)

        # 不处理空表格
        if len(headers) == 0:
            return None

        html_table += "  </tbody>\n</table>\n</body>\n</html>\n"

        # 在这里调用单元格合并方法
        # html_table = merge_cells_in_html(html_table)

        output_file = get_output_path(output_file_ext='html', output_filename=f'table_{page_number}_{self.table_num}', output_dir=f"{self.output_dir}/tables")
        with open(output_file, "w", encoding="utf-8") as file:
            file.write(html_table)
            self.table_num += 1
        return LoaderTable(page_number, doc_filename, title_stack, headers, html_table, level_stack)

    def get_cell_text(self, cell, doc_filename, level_stack, page_number, title_stack):
        """
        将表格cell中的文本和图片抽取
        :param cell:
        :param doc_filename:
        :param level_stack:
        :param page_number:
        :param title_stack:
        :return:
        """
        cell_text = cell.text.strip()
        # TODO 表格中的图片不处理
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run._r.xpath('.//pic:pic'):
                    for img in run._r.xpath('.//pic:pic'):
                        # 获取图片信息
                        image_info = self.get_image_info(img, page_number,
                                                         title_stack,
                                                         level_stack, doc_filename, None,
                                                         f"table_{page_number}")
                        # if image_info:
                        #     absolute_path = os.path.abspath(image_info.url)
                            # cell_text += f'<img src="{absolute_path}" alt="Image" />'
        return cell_text

    def get_image_info(self, image, page_number, title_stack, level_stack, doc_filename, surrounding_text,
                       prefix='img') -> LoaderImage:
        """获取图片信息"""
        """解析图片的宽度和高度"""
        embed = image.xpath('.//a:blip/@r:embed')[0]
        related_part = self.doc.part.related_parts[embed]
        # 已经处理过的图片不再做处理
        if embed in self.img_ids:
            return None
        self.img_ids.add(embed)
        image = related_part.image
        # 检查图片尺寸
        if image.px_width < self.min_img_width or image.px_height < self.min_img_height:
            return None
        image = Image.open(BytesIO(image.blob))
        file_path = get_output_path(output_file_ext='png', output_filename=f'{prefix}_{self.image_num}', output_dir=f"{self.output_dir}/images")
        # TODO 图片存储
        image_url = self.save_image(image, file_path)
        self.image_num += 1
        # 获取图片的宽度和高度
        return LoaderImage(page_number, doc_filename, title_stack, level_stack,
                           image.width, image.height,
                           None, image_url)

    def find_previous_text_paragraph(self, paragraphs, idx):
        """从给定的段落列表中向前查找最近的包含文本的段落"""
        while idx >= 0:
            if paragraphs[idx].text.strip():
                return paragraphs[idx].text
            idx -= 1
        return ""

    def get_element_from_xml(self, element, doc):
        if isinstance(element, CT_P):
            return Paragraph(element, doc)
        elif isinstance(element, CT_Tbl):
            return Table(element, doc)

    def extract_tables_from_paragraph(self, paragraph):
        tables = []
        for run in paragraph.runs:
            tbls = run._r.xpath('.//w:tbl')
            for tbl in tbls:
                tables.append(tbl.xml)  # 注意这里返回的是 XML 字符串
        return tables

    def create_table_from_xml(self, tbl_xml, parent_element):
        # 将 XML 字符串转换为 Open XML 元素
        tbl = parse_xml(tbl_xml)
        # 创建 Table 对象
        return Table(tbl, parent_element)  # 注意这里的 document 参数

    def load_and_parse(self, from_page=0, to_page=100000, callback=None):
        json_file_path = get_output_path(output_filename='result', output_file_ext='json',
                                         output_dir=f"{self.output_dir}/json")

        # 遍历文档的_body属性来处理段落、表格和图片
        body_elements = iter(self.doc.element.body)
        last_paragraph_text = ''
        for i, element in enumerate(body_elements):
            if self.page_number > to_page:
                break
            if from_page <= self.page_number < to_page:
                element = self.get_element_from_xml(element, self.doc)
                if isinstance(element, Paragraph):
                    p = element
                    # 收集toc条目
                    if self.is_toc_entry(p.style.name):
                        # 清洗文本，去除特殊符号和前缀
                        cleaned_text = self.clean_text(p.text)
                        level = int(re.findall(r'\d+', p.style.name)[-1])
                        toc_header = {
                            "header": cleaned_text,
                            "level": level
                        }
                        self.toc_max_length = max(len(cleaned_text), self.toc_max_length)
                        self.toc_entries.append(toc_header)
                        continue
                    # 判断是否为标题
                    title_level, p_text, title_type, is_title = self.docx_title_find(p)
                    if is_title:  # is a heading or matches toc entry
                        # 创建一个新的段落信息
                        if last_paragraph_text.strip():
                            paragraph_text = last_paragraph_text
                            # TODO 优化为检测stack中的标题是否是弱标题
                            if title_type == 'weak':
                                paragraph_text = last_paragraph_text
                            ti_list_entry = self.get_paragraph_info(paragraph_text, self.page_number, self.title_stack,
                                                                    self.level_stack, self.filename)
                            self.output.append(ti_list_entry)
                        # 避免影响已经组织好的title
                        self.title_stack = self.title_stack.copy()
                        self.level_stack = self.level_stack.copy()
                        while self.title_stack and title_level <= self.level_stack[-1]:
                            self.title_stack.pop()
                            self.level_stack.pop()
                        self.title_stack.append(p_text)
                        self.level_stack.append(title_level)
                        last_paragraph_text = ''
                        has_img = any(run._r.xpath('.//pic:pic') for run in element.runs)
                        has_table = any(run._r.xpath('.//w:tbl') for run in element.runs)
                        #  一定要先处理表格，因为表格中可能嵌套图片，不先处理可能会把图片从表格中剥离
                        if has_table:
                            tables = self.extract_tables_from_paragraph(element)
                            for tbl in tables:
                                # 将提取出的表格转换为 Table 对象
                                extracted_table = self.create_table_from_xml(tbl, element)
                                # 处理表格
                                tbl_info = self.process_table_with_context(extracted_table, last_paragraph_text,
                                                                           self.output,
                                                                           self.page_number, self.title_stack,
                                                                           self.level_stack,
                                                                           self.filename)
                                if tbl_info:
                                    self.output.append(tbl_info)
                            # print("找到了错误布局的表格")
                        # 标题所在的段落也可能有图片
                        if has_img:
                            # 处理段落内的图片
                            _, image_info = self.process_paragraph_with_image(p, last_paragraph_text, self.output,
                                                                              self.page_number, self.title_stack,
                                                                              self.level_stack, self.filename)
                            if image_info:
                                self.output.extend(image_info)

                    else:  # not a question or does not match toc entry
                        # 创建一个新的段落信息
                        last_paragraph_text += p.text.strip()
                        has_img = any(run._r.xpath('.//pic:pic') for run in element.runs)
                        has_table = any(run._r.xpath('.//w:tbl') for run in element.runs)
                        if (has_img or has_table) and last_paragraph_text:
                            ti_list_entry = self.get_paragraph_info(last_paragraph_text, self.page_number,
                                                                    self.title_stack,
                                                                    self.level_stack, self.filename)
                            self.output.append(ti_list_entry)
                            last_paragraph_text = ''
                        #  一定要先处理表格，因为表格中可能嵌套图片，不先处理可能会把图片从表格中剥离
                        if has_table:
                            tables = self.extract_tables_from_paragraph(element)
                            for tbl in tables:
                                # 将提取出的表格转换为 Table 对象
                                extracted_table = self.create_table_from_xml(tbl, element)
                                # 处理表格
                                tbl_info = self.process_table_with_context(extracted_table, last_paragraph_text,
                                                                           self.output,
                                                                           self.page_number, self.title_stack,
                                                                           self.level_stack,
                                                                           self.filename)
                                if tbl_info:
                                    self.output.append(tbl_info)
                        if has_img:
                            # 处理段落内的图片
                            _, image_info = self.process_paragraph_with_image(p, last_paragraph_text, self.output,
                                                                              self.page_number, self.title_stack,
                                                                              self.level_stack, self.filename)
                            if image_info:
                                self.output.extend(image_info)

                        # 创建一个新的段落信息
                elif isinstance(element, Table):

                    # 直接使用当前的title_stack和level_stack
                    tbl_info = self.process_table_with_context(element, last_paragraph_text, self.output,
                                                               self.page_number, self.title_stack, self.level_stack,
                                                               self.filename)
                    if tbl_info and last_paragraph_text:
                        ti_list_entry = self.get_paragraph_info(last_paragraph_text, self.page_number,
                                                                self.title_stack,
                                                                self.level_stack, self.filename)
                        self.output.append(ti_list_entry)
                        last_paragraph_text = ''
                    if tbl_info:
                        self.output.append(tbl_info)

            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    self.page_number += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    self.page_number += 1
        # 最后没有标题切换了，组织剩下的文本
        if last_paragraph_text:
            ti_list_entry = self.get_paragraph_info(last_paragraph_text, self.page_number, self.title_stack,
                                                    self.level_stack, self.filename)
            self.output.append(ti_list_entry)

        # 保存为 JSON 文件
        with open(json_file_path, "w", encoding="utf-8") as json_file:
            json.dump(self.output, json_file, cls=LoaderOutPutEncoder, ensure_ascii=False, indent=4)
        return self.output


if __name__ == '__main__':
    filename = f"docs/新能源汽车发展政府规划.docx"
    docx = DocxLoader(filename=filename)

    docx.load_and_parse()
